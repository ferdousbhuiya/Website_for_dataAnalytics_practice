from docx import Document

doc = Document(r"g:\Website_for_dataAnalytics_practice\statistics.docx")

# Extract all paragraphs
content = []
for para in doc.paragraphs:
    if para.text.strip():
        content.append(para.text)

# Extract tables
for table in doc.tables:
    content.append("\n[TABLE START]")
    for row in table.rows:
        row_data = [cell.text for cell in row.cells]
        content.append(" | ".join(row_data))
    content.append("[TABLE END]")

# Print full content
full_content = "\n".join(content)
# Save to file to handle encoding
with open(r"g:\Website_for_dataAnalytics_practice\extracted_stats.txt", "w", encoding="utf-8") as f:
    f.write(full_content)
print(f"Extracted {len(content)} paragraphs/rows. Saved to extracted_stats.txt")
